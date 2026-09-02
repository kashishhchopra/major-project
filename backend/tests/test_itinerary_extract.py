"""Itinerary document text extraction + heuristic parsing
(services/itinerary_extract.py)."""
import io

import pytest

from app.services.itinerary_extract import ExtractionError, extract_text, parse_itinerary_text


# ---------------------------------------------------------------- extraction
def test_extract_text_plain_text():
    text = extract_text("trip.txt", b"Delhi -> Agra -> Jaipur", "text/plain")
    assert "Delhi" in text


def test_extract_text_pdf():
    import pypdf

    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    # A blank PDF has no extractable text -- extraction should raise, not
    # silently return an empty itinerary.
    with pytest.raises(ExtractionError, match="scanned image|no extractable"):
        extract_text("trip.pdf", buf.getvalue(), "application/pdf")


def test_extract_text_docx():
    import docx

    document = docx.Document()
    document.add_paragraph("Day 1: Delhi")
    document.add_paragraph("Day 2: Agra")
    buf = io.BytesIO()
    document.save(buf)
    text = extract_text(
        "trip.docx", buf.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert "Delhi" in text and "Agra" in text


def _make_image_bytes(text: str, fmt: str = "JPEG") -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (700, 150), "white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 10), text, fill="black", font=ImageFont.load_default(size=40))
    buf = io.BytesIO()
    img.save(buf, format=fmt)
    return buf.getvalue()


def test_extract_text_image_ocr_reads_real_text():
    # Real Tesseract OCR (no API key needed) on a generated photo of an
    # itinerary line -- this is the actual jpg/png upload path, not a mock.
    content = _make_image_bytes("Delhi -> Agra")
    text = extract_text("scan.jpg", content, "image/jpeg")
    assert "Delhi" in text and "Agra" in text


def test_extract_text_image_png_also_supported():
    content = _make_image_bytes("Day 1: Jaipur", fmt="PNG")
    text = extract_text("scan.png", content, "image/png")
    assert "Jaipur" in text


def test_extract_text_image_with_no_readable_text_raises():
    from PIL import Image

    blank = Image.new("RGB", (100, 100), "white")
    buf = io.BytesIO()
    blank.save(buf, format="JPEG")
    with pytest.raises(ExtractionError, match="No readable text"):
        extract_text("blank.jpg", buf.getvalue(), "image/jpeg")


def test_extract_text_image_corrupt_data_raises_clear_error():
    with pytest.raises(ExtractionError, match="Could not read this image"):
        extract_text("scan.jpg", b"\xff\xd8\xff-not-a-real-image", "image/jpeg")


def test_extract_text_image_falls_back_gracefully_when_ocr_unavailable(monkeypatch):
    # Simulate a deployment with no `tesseract` binary installed -- the
    # feature must degrade to a clear message, never a 500 or a crash.
    import pytesseract

    def _raise(*a, **k):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(pytesseract, "image_to_string", _raise)
    content = _make_image_bytes("Delhi -> Agra")
    with pytest.raises(ExtractionError, match="OCR"):
        extract_text("scan.jpg", content, "image/jpeg")


def test_extract_text_rejects_unsupported_type():
    with pytest.raises(ExtractionError, match="Unsupported"):
        extract_text("trip.xyz", b"data", "application/x-unknown")


# ---------------------------------------------------------------- parsing
def test_parse_arrow_destination_sequence():
    result = parse_itinerary_text("Delhi -> Agra -> Jaipur -> Udaipur")
    names = [d["name"] for d in result["destinations"]]
    assert names == ["Delhi", "Agra", "Jaipur", "Udaipur"]


def test_parse_day_headers():
    result = parse_itinerary_text("Day 1: Delhi\nDay 2: Agra\nDay 3: Jaipur")
    names = [d["name"] for d in result["destinations"]]
    assert names == ["Delhi", "Agra", "Jaipur"]


def test_parse_ignores_garbled_day_header_noise():
    # OCR on a real photo sometimes reads several "Day N" headers onto one
    # line (columns/labels running together) -- that's noise, not a
    # destination, and must never be handed to the geocoder.
    result = parse_itinerary_text("Day 4 Day3 Day 2\nDay 1: Delhi")
    names = [d["name"] for d in result["destinations"]]
    assert names == ["Delhi"]
    assert any("day" in a.lower() for a in result["activities"])


def test_parse_finds_known_city_coordinates():
    result = parse_itinerary_text("Delhi -> Agra")
    delhi = result["destinations"][0]
    assert delhi["lat"] is not None and delhi["lng"] is not None
    assert delhi["location_demo"] is True  # gazetteer fallback, no Maps API key


def test_parse_hotel_lines():
    result = parse_itinerary_text("Taj Hotel Agra - Check-in 6 PM\nSome Resort - Check-out 10 AM")
    assert len(result["hotels"]) == 2
    assert result["hotels"][0]["check_in"] is True
    assert result["hotels"][1]["check_out"] is True


def test_parse_transport_lines():
    result = parse_itinerary_text("Flight AI-101 Delhi to Agra\nTrain 12002 Agra to Jaipur")
    assert len(result["transport"]) == 2


def test_parse_extracts_dates():
    result = parse_itinerary_text("Trip: 01 Sep 2026 to 10 Sep 2026\nDelhi -> Agra")
    assert result["trip_start"] == "2026-09-01"
    assert result["trip_end"] == "2026-09-10"


def test_parse_leftover_lines_become_activities():
    result = parse_itinerary_text("Visit the Taj Mahal at sunrise")
    assert "Visit the Taj Mahal at sunrise" in result["activities"]


def test_parse_never_raises_on_garbage_input():
    result = parse_itinerary_text("")
    assert result["destinations"] == []
    assert result["trip_start"] is None
